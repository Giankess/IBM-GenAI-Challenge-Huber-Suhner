from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import os
import tempfile
import shutil
from docx import Document
from docx.shared import RGB
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import torch
import uuid
import json
import datetime
from train import train_model, create_dataset, activate_model
from redline_parser import process_redline_document

app = FastAPI(title="NDA Validator API")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create directories for storing documents and memory
os.makedirs("documents", exist_ok=True)
os.makedirs("memory", exist_ok=True)
os.makedirs("training_data", exist_ok=True)
os.makedirs("models", exist_ok=True)

# Initialize versions.json if it doesn't exist
if not os.path.exists("models/versions.json"):
    with open("models/versions.json", "w") as f:
        json.dump([], f)

# Initialize the legal-bert model
tokenizer = AutoTokenizer.from_pretrained("nlpaueb/legal-bert-base-uncased")
model = AutoModelForSequenceClassification.from_pretrained("nlpaueb/legal-bert-base-uncased", num_labels=2)

# Memory storage
memory_file = "memory/memory.json"
if not os.path.exists(memory_file):
    with open(memory_file, "w") as f:
        json.dump([], f)

# Load replacements database if it exists
replacements_db = {}
for dataset_dir in os.listdir("training_data"):
    replacements_path = f"training_data/{dataset_dir}/replacements.json"
    if os.path.exists(replacements_path):
        try:
            with open(replacements_path, 'r') as f:
                replacements = json.load(f)
                for item in replacements:
                    if "original" in item and "replacement" in item:
                        replacements_db[item["original"]] = item["replacement"]
        except Exception as e:
            print(f"Error loading replacements from {replacements_path}: {str(e)}")

class FeedbackModel(BaseModel):
    document_id: str
    feedback: str

class DatasetModel(BaseModel):
    name: str
    is_redline: bool = False

class TrainingModel(BaseModel):
    dataset_id: str
    epochs: int = 3
    batch_size: int = 8
    learning_rate: float = 2e-5

class ModelActivationModel(BaseModel):
    model_version_id: str

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """
    Upload and process a document
    """
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    # Generate unique document ID
    document_id = str(uuid.uuid4())
    print(f"Processing document with ID: {document_id}")
    
    # Save the uploaded file
    file_path = f"documents/{document_id}.docx"
    try:
        print(f"Saving uploaded file to: {file_path}")
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        # Verify file was saved
        if not os.path.exists(file_path):
            raise HTTPException(status_code=500, detail="Failed to save uploaded file")
        
        file_size = os.path.getsize(file_path)
        print(f"Successfully saved file: {file_path} (size: {file_size} bytes)")
        
        # Process the document
        try:
            # Load the document
            doc = Document(file_path)
            
            # Parse document content
            doc_content = parse_document(doc)
            print(f"Parsed {len(doc_content)} paragraphs from document")
            
            # Check for problematic clauses
            problematic_clauses = check_document(doc_content)
            print(f"Found {len(problematic_clauses)} problematic clauses")
            
            # Generate suggestions
            suggestions = make_suggestions(problematic_clauses)
            print(f"Generated {len(suggestions)} suggestions")
            
            # Create redline document
            redline_path = create_redline_document(document_id, file_path, suggestions)
            print(f"Created redline document at: {redline_path}")
            
            # Verify redline document was created
            if not os.path.exists(redline_path):
                raise HTTPException(status_code=500, detail="Failed to create redline document")
            
            # Save to memory
            save_to_memory(document_id, {
                "original_file": file_path,
                "redline_file": redline_path,
                "suggestions": suggestions,
                "original_filename": file.filename.replace('.docx', '')
            })
            print("Successfully saved document data to memory")
            
            return {
                "document_id": document_id,
                "suggestions": suggestions
            }
            
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            # Clean up files if processing fails
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
            
    except Exception as e:
        print(f"Error saving uploaded file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error saving uploaded file: {str(e)}")

def parse_document(doc):
    """
    Parse the Word document and extract text
    """
    content = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            content.append({
                "text": para.text,
                "index": len(content)
            })
    
    return content

def check_document(doc_content):
    """
    Check the document for problematic clauses using the legal-bert model
    """
    problematic_clauses = []
    
    # Initialize model and tokenizer
    model = AutoModelForSequenceClassification.from_pretrained("nlpaueb/legal-bert-base-uncased", num_labels=2)
    tokenizer = AutoTokenizer.from_pretrained("nlpaueb/legal-bert-base-uncased")
    
    # Check if we have an active fine-tuned model
    model_path = "models/active_model"
    if os.path.exists(model_path) and os.path.isdir(model_path):
        # Load the fine-tuned model
        try:
            model = AutoModelForSequenceClassification.from_pretrained(model_path)
            tokenizer = AutoTokenizer.from_pretrained(model_path)
        except Exception as e:
            print(f"Error loading fine-tuned model: {str(e)}")
            # Fall back to the base model (already initialized above)
            pass
    
    # Keywords that indicate non-legal content
    non_legal_keywords = [
        "address", "street", "city", "country", "postal", "zip",
        "name", "mr.", "mrs.", "ms.", "dr.", "prof.", "company",
        "inc.", "ltd.", "llc", "gmbh", "ag", "sa", "sarl",
        "date:", "dated:", "signature", "signed", "witness"
    ]
    
    for clause in doc_content:
        text = clause["text"].lower()
        
        # Skip if the clause contains non-legal content
        if any(keyword in text for keyword in non_legal_keywords):
            continue
            
        # Skip very short clauses (likely not legal content)
        if len(text.split()) < 5:
            continue
            
        # Tokenize the text
        inputs = tokenizer(clause["text"], return_tensors="pt", truncation=True, max_length=512)
        
        # Get model prediction
        with torch.no_grad():
            outputs = model(**inputs)
            prediction = torch.softmax(outputs.logits, dim=1)
            
        # If the model predicts this is a problematic clause (class 1)
        if prediction[0][1] > 0.5:
            problematic_clauses.append({
                "index": clause["index"],
                "text": clause["text"],
                "score": float(prediction[0][1])
            })
    
    return problematic_clauses

def make_suggestions(problematic_clauses):
    """
    Generate suggestions for problematic clauses
    """
    suggestions = []
    
    for clause in problematic_clauses:
        # Check if we have a replacement for this clause in our database
        replacement_text = None
        if clause["text"] in replacements_db:
            replacement_text = replacements_db[clause["text"]]
        
        # If no exact match, try to find a similar clause
        if not replacement_text:
            for original, replacement in replacements_db.items():
                # Simple similarity check - in a real implementation, use a more sophisticated approach
                if len(original) > 20 and original in clause["text"]:
                    replacement_text = replacement
                    break
        
        # If we found a replacement, use it
        if replacement_text:
            suggestion = {
                "index": clause["index"],
                "original": clause["text"],
                "suggested": f"SUGGESTED CHANGE: {replacement_text}",
                "reason": "This clause has been identified as problematic based on previous redlines."
            }
        else:
            # Otherwise use a generic suggestion
            suggestion = {
                "index": clause["index"],
                "original": clause["text"],
                "suggested": f"SUGGESTED CHANGE: {clause['text']} [This clause has been identified as potentially problematic and should be reviewed]",
                "reason": "This clause may contain terms that are unfavorable or legally problematic."
            }
        
        suggestions.append(suggestion)
    
    return suggestions

def create_redline_document(document_id, original_path, suggestions):
    """
    Create a redline document with the suggestions
    """
    doc = Document(original_path)
    
    # Map suggestions to paragraph indices
    suggestion_map = {s["index"]: s for s in suggestions}
    
    # Apply suggestions to the document
    for i, para in enumerate(doc.paragraphs):
        if i in suggestion_map:
            # Clear existing runs
            for run in para.runs:
                run.text = ""
            
            # Add original text in red strikethrough
            original_run = para.add_run(suggestion_map[i]["original"])
            original_run.font.strike = True
            original_run.font.color.rgb = RGB(255, 0, 0)  # Red color
            
            # Add new line
            para.add_run("\n")
            
            # Add suggested text in green
            suggested_text = suggestion_map[i]["suggested"]
            if "SUGGESTED CHANGE: " in suggested_text:
                suggested_text = suggested_text.split("SUGGESTED CHANGE: ")[1]
            suggested_run = para.add_run(suggested_text)
            suggested_run.font.color.rgb = RGB(0, 128, 0)  # Green color
    
    # Save the redline document
    redline_path = f"documents/{document_id}_redline.docx"
    doc.save(redline_path)
    
    return redline_path

def save_to_memory(document_id, data):
    """
    Save document data to memory
    """
    with open(memory_file, "r") as f:
        memory = json.load(f)
    
    memory.append({
        "document_id": document_id,
        "data": data,
        "timestamp": str(datetime.datetime.now())
    })
    
    with open(memory_file, "w") as f:
        json.dump(memory, f)

@app.get("/download/{document_id}/{type}")
async def download_document(document_id: str, type: str):
    """
    Download a document (redline or clean)
    """
    if type not in ["redline", "clean"]:
        raise HTTPException(status_code=400, detail="Invalid document type")
    
    file_path = f"documents/{document_id}_{type}.docx"
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Document not found")
    
    return FileResponse(file_path, filename=f"nda_{type}.docx")

@app.post("/feedback")
async def submit_feedback(feedback: FeedbackModel):
    """
    Submit feedback on the redline document
    """
    # Get document data from memory
    with open(memory_file, "r") as f:
        memory = json.load(f)
    
    document_data = None
    for item in memory:
        if item["document_id"] == feedback.document_id:
            document_data = item["data"]
            break
    
    if not document_data:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Interpret feedback
    updated_suggestions = interpret_feedback(feedback.feedback, document_data["suggestions"])
    
    # Create new redline document
    redline_path = create_redline_document(
        feedback.document_id, 
        document_data["original_file"], 
        updated_suggestions
    )
    
    # Update memory
    for item in memory:
        if item["document_id"] == feedback.document_id:
            item["data"]["suggestions"] = updated_suggestions
            item["data"]["redline_file"] = redline_path
            item["data"]["feedback"] = feedback.feedback
            break
    
    with open(memory_file, "w") as f:
        json.dump(memory, f)
    
    return {"document_id": feedback.document_id, "redline_file": redline_path}

def interpret_feedback(feedback_text, suggestions):
    """
    Interpret user feedback and update suggestions
    """
    # In a real implementation, this would use NLP to understand the feedback
    # For now, we'll just make a simple modification to the suggestions
    updated_suggestions = []
    
    for suggestion in suggestions:
        updated_suggestion = suggestion.copy()
        updated_suggestion["suggested"] = f"UPDATED BASED ON FEEDBACK: {suggestion['suggested']}"
        updated_suggestions.append(updated_suggestion)
    
    return updated_suggestions

@app.post("/accept/{document_id}")
async def accept_suggestions(document_id: str):
    """
    Accept all suggestions and create a clean document
    """
    # Get document data from memory
    with open(memory_file, "r") as f:
        memory = json.load(f)
    
    document_data = None
    for item in memory:
        if item["document_id"] == document_id:
            document_data = item["data"]
            break
    
    if not document_data:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Create clean document
    clean_path = create_clean_document(document_id, document_data["original_file"], document_data["suggestions"])
    
    # Update memory
    for item in memory:
        if item["document_id"] == document_id:
            item["data"]["clean_file"] = clean_path
            break
    
    with open(memory_file, "w") as f:
        json.dump(memory, f)
    
    return {"document_id": document_id, "clean_file": clean_path}

def create_clean_document(document_id, original_path, suggestions):
    """
    Create a clean document with all suggestions accepted
    """
    doc = Document(original_path)
    
    # Map suggestions to paragraph indices
    suggestion_map = {s["index"]: s for s in suggestions}
    
    # Apply suggestions to the document (without redlining)
    for i, para in enumerate(doc.paragraphs):
        if i in suggestion_map:
            # Extract just the suggested text without redline formatting
            suggested_text = suggestion_map[i]["suggested"]
            if "SUGGESTED CHANGE: " in suggested_text:
                suggested_text = suggested_text.split("SUGGESTED CHANGE: ")[1]
                if " [This clause" in suggested_text:
                    suggested_text = suggested_text.split(" [This clause")[0]
            
            para.text = suggested_text
    
    # Save the clean document
    clean_path = f"documents/{document_id}_clean.docx"
    doc.save(clean_path)
    
    return clean_path

# Training API endpoints

@app.post("/datasets")
async def create_training_dataset(name: str = Form(...), is_redline: bool = Form(False), files: List[UploadFile] = File(...)):
    """
    Create a new training dataset from uploaded files
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    for file in files:
        if not file.filename.endswith(('.docx', '.doc', '.json')):
            raise HTTPException(status_code=400, detail="Only Word documents and JSON annotation files are supported")
    
    dataset_id, metadata = create_dataset(name, files, is_redline)
    
    return {
        "dataset_id": dataset_id,
        "name": metadata["name"],
        "document_count": metadata["document_count"],
        "created_at": metadata["created_at"],
        "is_redline": metadata.get("is_redline", False)
    }

@app.get("/datasets")
async def list_datasets():
    """
    List all available training datasets
    """
    datasets = []
    
    for dataset_dir in os.listdir("training_data"):
        metadata_path = f"training_data/{dataset_dir}/metadata.json"
        if os.path.exists(metadata_path):
            with open(metadata_path, 'r') as f:
                metadata = json.load(f)
                datasets.append(metadata)
    
    return datasets

@app.post("/training")
async def start_training(training: TrainingModel, background_tasks: BackgroundTasks):
    """
    Start a training job for a dataset
    """
    # Verify the dataset exists
    dataset_path = f"training_data/{training.dataset_id}"
    if not os.path.exists(dataset_path):
        raise HTTPException(status_code=404, detail="Dataset not found")
    
    # Start training in the background
    background_tasks.add_task(
        train_model,
        training.dataset_id,
        training.epochs,
        training.batch_size,
        training.learning_rate
    )
    
    return {
        "message": "Training job started",
        "dataset_id": training.dataset_id
    }

@app.get("/training")
async def list_training_jobs():
    """
    List all training jobs
    """
    jobs = []
    
    for job_dir in os.listdir("models"):
        if job_dir == "versions.json" or job_dir == "active_model":
            continue
            
        metadata_path = f"models/{job_dir}/metadata.json"
        if os.path.exists(metadata_path):
            with open(metadata_path, 'r') as f:
                metadata = json.load(f)
                jobs.append(metadata)
    
    return jobs

@app.get("/models")
async def list_model_versions():
    """
    List all model versions
    """
    with open("models/versions.json", 'r') as f:
        try:
            versions = json.load(f)
        except json.JSONDecodeError:
            versions = []
    
    return versions

@app.post("/models/activate")
async def activate_model_version(activation: ModelActivationModel):
    """
    Activate a model version for use in the NDA validator
    """
    success = activate_model(activation.model_version_id)
    
    if not success:
        raise HTTPException(status_code=500, detail="Failed to activate model")
    
    return {
        "message": "Model activated successfully",
        "model_version_id": activation.model_version_id
    }

@app.post("/parse-redline")
async def parse_redline(file: UploadFile = File(...)):
    """
    Parse a redline document and return the extracted problematic clauses and replacements
    """
    if not file.filename.endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Only Word documents are supported")
    
    # Save the uploaded file temporarily
    temp_file = f"temp_{uuid.uuid4()}.docx"
    with open(temp_file, "wb") as f:
        shutil.copyfileobj(file.file, f)
    
    try:
        # Process the redline document
        training_data = process_redline_document(temp_file)
        
        # Count problematic clauses
        problematic_count = sum(1 for clause in training_data["clauses"] if clause.get("is_problematic", False))
        
        return {
            "total_clauses": len(training_data["clauses"]),
            "problematic_clauses": problematic_count,
            "data": training_data
        }
    finally:
        # Clean up the temporary file
        if os.path.exists(temp_file):
            os.remove(temp_file)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
